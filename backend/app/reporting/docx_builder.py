"""DOCX report renderer using python-docx."""
from __future__ import annotations
from datetime import datetime, timezone
from pathlib import Path

def render_docx_report(eng, findings, report, out_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        # Fallback: write a placeholder
        out_path.write_bytes(b"python-docx not installed")
        return

    doc = Document()
    doc.add_heading(report.title, 0)
    eng_name = eng.name if eng else "Unknown Engagement"
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Engagement: {eng_name}")
    doc.add_paragraph(f"Generated: {now}")
    if report.is_draft:
        p = doc.add_paragraph("DRAFT – Not for Distribution")
        p.runs[0].bold = True

    doc.add_heading("Finding Summary", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Title"
    hdr[2].text = "Severity"
    hdr[3].text = "Status"
    for i, f in enumerate(findings, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = f.title
        row[2].text = f.severity.upper()
        row[3].text = f.status

    doc.add_heading("Detailed Findings", 1)
    for i, f in enumerate(findings, 1):
        doc.add_heading(f"F{i:03d} – {f.title}", 2)
        doc.add_paragraph(f"Severity: {f.severity.upper()} | Status: {f.status}")
        if f.description:
            doc.add_heading("Description", 3)
            doc.add_paragraph(f.description)
        if f.impact:
            doc.add_heading("Impact", 3)
            doc.add_paragraph(f.impact)
        if f.remediation:
            doc.add_heading("Remediation", 3)
            doc.add_paragraph(f.remediation)

    doc.save(str(out_path))
